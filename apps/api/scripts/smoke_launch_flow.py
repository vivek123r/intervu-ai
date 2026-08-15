from __future__ import annotations

import argparse
import asyncio
import io
import json
from datetime import UTC, datetime, timedelta
from typing import Any

import httpx
from docx import Document


class SmokeFailure(RuntimeError):
    pass


def build_resume() -> bytes:
    document = Document()
    document.add_heading("Aarav Mehta", level=1)
    document.add_paragraph("Backend Engineer — Bengaluru, India")
    document.add_heading("Experience", level=2)
    document.add_paragraph(
        "Backend Engineer at Northstar Systems. Designed Node.js and Python services, "
        "PostgreSQL transaction workflows, Redis caching, and event-driven APIs."
    )
    document.add_paragraph(
        "Reduced checkout latency by 34%, introduced idempotency keys for payment writes, "
        "and built dashboards for lock contention, retries, and error budgets."
    )
    document.add_heading("Projects", level=2)
    document.add_paragraph(
        "Built a resilient order platform with optimistic concurrency control, transactional "
        "outbox delivery, Docker, AWS, and OpenTelemetry."
    )
    document.add_heading("Skills", level=2)
    document.add_paragraph(
        "Python, FastAPI, Node.js, TypeScript, PostgreSQL, SQL, Redis, REST APIs, Docker, AWS, "
        "system design, observability, distributed systems"
    )
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class LaunchFlow:
    def __init__(self, base_url: str, timeout: float) -> None:
        self.client = httpx.AsyncClient(
            base_url=base_url.rstrip("/"),
            headers={"Authorization": "Bearer demo-token"},
            timeout=timeout,
        )
        self.phases: list[str] = []

    async def close(self) -> None:
        await self.client.aclose()

    async def request(
        self,
        method: str,
        path: str,
        *,
        phase: str,
        expected: int = 200,
        **kwargs: Any,
    ) -> dict[str, Any]:
        response = await self.client.request(method, path, **kwargs)
        if response.status_code != expected:
            try:
                payload = response.json()
                error = payload.get("error", {})
                reason = error.get("code") or error.get("message") or "unknown error"
            except (json.JSONDecodeError, AttributeError):
                reason = "non-JSON response"
            raise SmokeFailure(f"{phase} failed ({response.status_code}): {reason}")
        self.phases.append(phase)
        if response.status_code == 204:
            return {}
        payload = response.json()
        if not isinstance(payload, dict):
            raise SmokeFailure(f"{phase} returned an unexpected response shape")
        return payload

    async def run(self) -> dict[str, Any]:
        await self.request("GET", "/api/v1/users/me", phase="identity")

        now = datetime.now(UTC)
        interview = await self.request(
            "POST",
            "/api/v1/interviews",
            phase="interview_created",
            expected=201,
            json={
                "company_name": "Northstar Systems",
                "role_title": "Senior Backend Engineer",
                "interview_type": "technical",
                "round_name": "Architecture and Reliability",
                "round_number": 2,
                "total_rounds": 4,
                "scheduled_at": (now + timedelta(days=5)).isoformat(),
                "timezone": "Asia/Kolkata",
                "duration_minutes": 60,
                "meeting_type": "Google Meet",
                "rounds": [
                    {
                        "position": 1,
                        "name": "Recruiter Screen",
                        "type": "recruiter",
                        "status": "completed",
                    },
                    {
                        "position": 2,
                        "name": "Architecture and Reliability",
                        "type": "technical",
                        "status": "upcoming",
                    },
                ],
            },
        )
        interview_id = interview["id"]

        await self.request(
            "POST",
            "/api/v1/resumes",
            phase="resume_analyzed",
            expected=201,
            data={"interview_id": interview_id, "is_primary": "true"},
            files={
                "file": (
                    "aarav-mehta-resume.docx",
                    build_resume(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

        job_description = await self.request(
            "POST",
            "/api/v1/job-descriptions",
            phase="job_description_analyzed",
            expected=201,
            json={
                "interview_id": interview_id,
                "company_name": "Northstar Systems",
                "role_title": "Senior Backend Engineer",
                "raw_text": (
                    "Northstar Systems is hiring a Senior Backend Engineer to design reliable "
                    "payment and order services. Required skills include Python or Node.js, "
                    "PostgreSQL, SQL transaction isolation, REST APIs, distributed systems, "
                    "system design, observability, and production debugging. Candidates should "
                    "have five years of backend experience and be able to explain consistency, "
                    "idempotency, failure recovery, and database tradeoffs. Redis, Docker, AWS, "
                    "Kafka, and mentoring experience are preferred. The engineer will own design "
                    "reviews, reliability metrics, incident follow-up, and scalable API delivery."
                ),
            },
        )

        plan = await self.request(
            "POST",
            f"/api/v1/interviews/{interview_id}/prepare",
            phase="preparation_generated",
        )
        if not plan.get("tasks"):
            raise SmokeFailure("preparation_generated returned no tasks")

        session = await self.request(
            "POST",
            "/api/v1/practice/sessions",
            phase="session_created",
            expected=201,
            json={
                "interview_id": interview_id,
                "mode": "technical",
                "difficulty": "hard",
                "interviewer_style": "strict_technical_lead",
                "planned_duration": 20,
                "focus_areas": ["SQL Transactions", "Failure Recovery"],
            },
        )
        session_id = session["id"]

        started = await self.request(
            "POST",
            f"/api/v1/practice/sessions/{session_id}/start",
            phase="adaptive_question_generated",
        )
        question = started.get("current_question")
        if not isinstance(question, dict) or "id" not in question:
            raise SmokeFailure("adaptive_question_generated returned no current question")

        answer_started = datetime.now(UTC)
        answer = await self.request(
            "POST",
            f"/api/v1/practice/sessions/{session_id}/answers",
            phase="answer_evaluated_and_routed",
            json={
                "question_id": question["id"],
                "transcript": (
                    "Suppose an order transaction updates two rows and the process crashes before "
                    "the commit record is durable. PostgreSQL writes changed-page records to its "
                    "write-ahead log before data pages, so recovery starts from the latest checkpoint, "
                    "replays committed WAL records, and leaves transactions without a durable commit "
                    "record invisible. That preserves atomicity and durability even if some dirty "
                    "pages reached disk. WAL records include enough ordering and page information for "
                    "redo, while MVCC prevents an uncommitted version from becoming visible. Recovery "
                    "replay is restartable and idempotent, so another crash during recovery resumes "
                    "from durable WAL and the checkpoint rather than treating partial replay as a "
                    "business commit. At the application boundary I would also use an idempotency key, "
                    "keep external calls outside the transaction, and publish committed work through "
                    "a transactional outbox. I would monitor checkpoint duration, WAL lag, failed "
                    "recovery, lock waits, duplicate attempts, and replica replay delay."
                ),
                "started_at": answer_started.isoformat(),
                "ended_at": (answer_started + timedelta(seconds=56)).isoformat(),
                "duration_ms": 56_000,
                "pause_markers_ms": [8_400, 23_500, 41_200],
            },
        )

        report = await self.request(
            "POST",
            f"/api/v1/practice/sessions/{session_id}/complete",
            phase="report_generated",
        )
        restored = await self.request(
            "GET",
            f"/api/v1/practice/sessions/{session_id}/report",
            phase="report_recovered",
        )
        analytics = await self.request(
            "GET",
            "/api/v1/analytics/overview",
            phase="analytics_updated",
        )

        if restored.get("id") != report.get("id"):
            raise SmokeFailure("report recovery returned a different report")
        speech_metrics = report.get("speech_metrics") or {}
        if not speech_metrics.get("total_words"):
            raise SmokeFailure("report is missing deterministic speech metrics")
        if float(report.get("overall_score") or 0) < 20:
            raise SmokeFailure("report score is implausibly low for the calibrated smoke answer")

        return {
            "status": "passed",
            "phases": self.phases,
            "role_match_score": (job_description.get("role_match_data") or {}).get("score"),
            "preparation_tasks": len(plan["tasks"]),
            "followup_decision": answer.get("decision"),
            "report_score": report.get("overall_score"),
            "speech_words": speech_metrics.get("total_words"),
            "analytics_questions_answered": analytics.get("questions_answered"),
        }


async def async_main() -> int:
    parser = argparse.ArgumentParser(description="Exercise the Intervu AI launch flow.")
    parser.add_argument("--base-url", default="http://localhost:8000")
    parser.add_argument("--timeout", type=float, default=120.0)
    args = parser.parse_args()

    flow = LaunchFlow(args.base_url, args.timeout)
    try:
        result = await flow.run()
    except (SmokeFailure, httpx.HTTPError) as exc:
        print(json.dumps({"status": "failed", "reason": str(exc)}, indent=2))
        return 1
    finally:
        await flow.close()
    print(json.dumps(result, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(asyncio.run(async_main()))
